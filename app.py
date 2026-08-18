from __future__ import annotations

import hashlib
import hmac
import json
import os
import secrets
import urllib.error
import urllib.request
import zipfile
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path

from google import genai
from google.genai import types as genai_types
from dotenv import load_dotenv
from flask import Flask, Response, abort, jsonify, make_response, render_template, request
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"
app = Flask(__name__)

# ── SEOページ関連の設定 ────────────────────────────────────────────────────────
# アプリ本体（/app）とは切り離した、検索エンジン向けの静的ページ群で使う定数・
# ページ一覧。sitemap.xml もこの一覧から自動生成する。

SITE_NAME = "まとめときや"
SITE_URL = "https://matome.webtool-labs.com"
FIREBASE_AUTH_DOMAIN = "memo-app-9dd98.firebaseapp.com"
OPERATOR_NAME = "RakuYade"
OPERATOR_PROFILE_URL = "https://profile.webtool-labs.com/"
OPERATOR_SITE_URL = "https://webtool-labs.com/"
CONTACT_FORM_URL = (
    "https://docs.google.com/forms/d/e/1FAIpQLSePfOxSZwYoGcL7csdt1RbLY4eQ9gdu6ePqcgZ96xTxZj8GXA/viewform?usp=publish-editor"
)

# 狙うキーワードごとの1ページ。テンプレートは templates/seo/keywords/ 配下。
KEYWORD_PAGES = [
    {
        "slug": "hierarchical-memo",
        "keyword": "階層メモ アプリ",
        "label": "階層メモアプリとは",
        "template": "seo/keywords/hierarchical-memo.html",
    },
    {
        "slug": "free-mindmap",
        "keyword": "マインドマップ 作成 無料",
        "label": "無料マインドマップ作成",
        "template": "seo/keywords/free-mindmap.html",
    },
    {
        "slug": "organize-ideas",
        "keyword": "メモ アイデア整理",
        "label": "アイデア整理のコツ",
        "template": "seo/keywords/organize-ideas.html",
    },
    {
        "slug": "ai-memo-generator",
        "keyword": "AI メモ 自動生成",
        "label": "AIでメモを自動生成",
        "template": "seo/keywords/ai-memo-generator.html",
    },
    {
        "slug": "collaborative-memo",
        "keyword": "共同編集 メモ帳",
        "label": "共同編集メモ帳",
        "template": "seo/keywords/collaborative-memo.html",
    },
    {
        "slug": "structured-notes",
        "keyword": "ノート 構造化",
        "label": "ノートを構造化する",
        "template": "seo/keywords/structured-notes.html",
    },
]
KEYWORD_PAGES_BY_SLUG = {page["slug"]: page for page in KEYWORD_PAGES}

# 悩み系キーワードで書くブログ記事。テンプレートは templates/seo/blog/ 配下。
BLOG_POSTS = [
    {
        "slug": "organize-scattered-ideas",
        "title": "アイデアが頭の中でごちゃごちゃになる人へ。階層メモで思考を整理するコツ",
        "template": "seo/blog/organize-scattered-ideas.html",
    },
    {
        "slug": "mindmap-vs-memo",
        "title": "マインドマップとメモ、結局どっちを使えばいい？使い分けの考え方",
        "template": "seo/blog/mindmap-vs-memo.html",
    },
    {
        "slug": "ai-brainstorming-tips",
        "title": "一人ブレストで手が止まったときの、AIとメモアプリの組み合わせ方",
        "template": "seo/blog/ai-brainstorming-tips.html",
    },
]
BLOG_POSTS_BY_SLUG = {post["slug"]: post for post in BLOG_POSTS}

# /faq とランディングページ抜粋、FAQPage構造化データで共用する質問一覧。
FAQ_ITEMS = [
    {
        "question": "無料で使えますか？",
        "answer": "はい、まとめときやは無料でご利用いただけます。メモ作成・階層管理・マインドマップ・AI生成・共同編集など、主要な機能はすべて追加費用なしで使えます。",
    },
    {
        "question": "会員登録は必要ですか？",
        "answer": "まとめときやを使うにはアカウント登録（メールアドレスまたはGoogleアカウント）が必要です。登録すればログインしたどの端末からでも続きから編集できます。",
    },
    {
        "question": "作成したメモはどこに保存されますか？",
        "answer": "メモはお使いのアカウントに紐づけてクラウド上に保存されます。端末を変えても、ログインすれば同じ内容を確認・編集できます。",
    },
    {
        "question": "他の人と一緒に編集できますか？",
        "answer": "はい。「共同編集」機能を使うと、合言葉を伝えるだけで複数人が同じメモをリアルタイムで編集できます。",
    },
    {
        "question": "AIでメモを作るとき、入力した内容は保存されますか？",
        "answer": "AI生成のために入力したテーマや添付ファイルは、生成処理のためだけに利用され、アップロード内容自体をサーバー側に保存することはありません。",
    },
    {
        "question": "作成したメモを他の形式で書き出せますか？",
        "answer": "メモはPDF・テキスト・Markdown形式で、マインドマップはPNG・SVG・PDF形式でダウンロードできます。",
    },
]


def get_public_pages() -> list[dict]:
    """sitemap.xml生成用の公開ページ一覧（パス・優先度・対応テンプレート）。
    アプリ内API・ログイン後にしか意味を持たないページは含めない。"""
    pages = [
        {"path": "/", "template": "seo/landing.html", "priority": "1.0"},
        {"path": "/app", "template": "index.html", "priority": "0.9"},
        {"path": "/how-to-use", "template": "seo/how_to_use.html", "priority": "0.6"},
        {"path": "/faq", "template": "seo/faq.html", "priority": "0.6"},
        {"path": "/about", "template": "seo/about.html", "priority": "0.4"},
        {"path": "/privacy", "template": "seo/privacy.html", "priority": "0.3"},
        {"path": "/terms", "template": "seo/terms.html", "priority": "0.3"},
        {"path": "/contact", "template": "seo/contact.html", "priority": "0.4"},
    ]
    for page in KEYWORD_PAGES:
        pages.append({"path": f"/{page['slug']}", "template": page["template"], "priority": "0.8"})
    for post in BLOG_POSTS:
        pages.append({"path": f"/blog/{post['slug']}", "template": post["template"], "priority": "0.6"})
    return pages


@app.context_processor
def inject_seo_globals():
    def static_version(filename: str) -> int:
        try:
            return int((BASE_DIR / "static" / filename).stat().st_mtime)
        except OSError:
            return 0

    return {
        "site_name": SITE_NAME,
        "site_url": SITE_URL,
        "operator_name": OPERATOR_NAME,
        "operator_profile_url": OPERATOR_PROFILE_URL,
        "operator_site_url": OPERATOR_SITE_URL,
        "contact_form_url": CONTACT_FORM_URL,
        "keyword_pages": KEYWORD_PAGES,
        "blog_posts": BLOG_POSTS,
        "canonical_url": f"{SITE_URL}{request.path}",
        "current_year": datetime.now(tz=timezone.utc).year,
        "static_version": static_version,
    }

_USE_VERTEX_AI = os.getenv("USE_VERTEX_AI", "false").lower() == "true"
_GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite")
_AI_FILE_MAX_BYTES = 10 * 1024 * 1024
_AI_FILE_TEXT_MAX_CHARS = 100_000
_AI_OFFICE_MAX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
_AI_DOCUMENT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".pdf", ".docx", ".xlsx"}
_AI_IMAGE_MIME_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".heic": "image/heic",
    ".heif": "image/heif",
}
_AI_FILE_EXTENSIONS = _AI_DOCUMENT_EXTENSIONS | set(_AI_IMAGE_MIME_TYPES)
_TODO_API_TOKEN_HASH_ENV = "MATOME_TODO_API_TOKEN_SHA256"
_TODO_API_WRITE_TOKEN_HASH_ENV = "MATOME_TODO_API_WRITE_TOKEN_SHA256"
_TODO_API_UID_ENV = "MATOME_TODO_API_UID"
_CLAUDE_PAIRING_CODE_LENGTH = 8
_CLAUDE_PAIRING_CODE_TTL = timedelta(minutes=10)
_firestore_client = None

app.config["MAX_CONTENT_LENGTH"] = _AI_FILE_MAX_BYTES + (512 * 1024)

_NOTE_PROMPT = """\
あなたは階層型メモ生成AIです。与えられたテーマをもとに、階層構造のメモをJSON形式で生成してください。

出力はJSONオブジェクトのみ（マークダウンコードブロックや説明文は不要）:
{{
  "title": "ルートメモのタイトル",
  "content": "ルートメモの内容（概要や導入文）",
  "children": [
    {{
      "title": "セクション見出し",
      "content": "このセクションの詳細な内容",
      "children": [
        {{
          "title": "サブ項目",
          "content": "サブ項目の内容",
          "children": []
        }}
      ]
    }}
  ]
}}

ルール:
- 入力と同じ言語で出力する
- ルート: テーマ全体のタイトルと概要
- セクション: 3〜5個のメインセクション
- 各セクション: 具体的な説明文を content に記載
- サブ項目: 必要に応じて1〜3個
- content は文章で書く（箇条書き不可）
- 出力: JSONオブジェクトのみ

テーマ：{prompt}"""

_MINDMAP_PROMPT = """\
あなたはマインドマップ生成AIです。与えられたテーマや内容をもとに、マインドマップをJSON形式で生成してください。

出力はJSONオブジェクトのみ（マークダウンコードブロックや説明文は不要）:
{{
  "title": "ルートノードのタイトル",
  "memo": "テーマ全体の要約、背景、重要な観点を1〜3文で説明する",
  "children": [
    {{
      "title": "ブランチ",
      "memo": "このブランチで扱う内容、なぜ重要か、具体的に考えるポイントを1〜3文で説明する",
      "children": [
        {{
          "title": "サブ項目",
          "memo": "このサブ項目の具体的な内容、例、実行のヒントを1〜2文で説明する",
          "children": []
        }}
      ]
    }}
  ]
}}

ルール:
- 入力と同じ言語で出力する（日本語入力→日本語ノード）
- すべてのノードで memo を必ず生成する（空文字は禁止）
- memo はノード本文として読める自然な文章にする
- 詳細説明・背景・具体例・注意点・次のアクションは title ではなく memo に書く
- ルートタイトル：簡潔に（15文字以内推奨）
- メインブランチ：3〜6個
- 各ブランチのサブ項目：1〜4個
- 最大深さ：ルート＋2レベル
- ノードタイトル：短く明確に（20文字以内推奨）
- 出力：JSONオブジェクトのみ

テーマ・内容：{prompt}"""


def create_gemini_client():
    api_key = os.getenv("GEMINI_API_KEY", "").strip()
    if _USE_VERTEX_AI:
        credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS", "").strip()
        if credentials_path and not Path(credentials_path).expanduser().is_file():
            if api_key:
                return genai.Client(api_key=api_key)
            raise RuntimeError(
                "Vertex AIの認証ファイルが見つかりません。"
                "GEMINI_API_KEYを設定するか、GOOGLE_APPLICATION_CREDENTIALSのパスを確認してください。"
            )
        project = os.getenv("GOOGLE_CLOUD_PROJECT", "")
        location = os.getenv("GOOGLE_CLOUD_LOCATION", "us-central1")
        if not project:
            if api_key:
                return genai.Client(api_key=api_key)
            raise RuntimeError("GOOGLE_CLOUD_PROJECT が設定されていません")
        return genai.Client(vertexai=True, project=project, location=location)

    if not api_key:
        raise RuntimeError("GEMINI_API_KEY が設定されていません")
    return genai.Client(api_key=api_key)


def _decode_text_file(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp932"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("ファイルの文字コードを読み取れませんでした。UTF-8形式で保存してください。")


def _validate_office_archive(data: bytes) -> None:
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            total_size = sum(info.file_size for info in archive.infolist())
    except zipfile.BadZipFile as exc:
        raise ValueError("Officeファイルが破損しているか、形式が正しくありません。") from exc
    if total_size > _AI_OFFICE_MAX_UNCOMPRESSED_BYTES:
        raise ValueError("展開後のファイルサイズが大きすぎます。")


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    if reader.is_encrypted:
        try:
            unlocked = reader.decrypt("")
        except Exception as exc:
            raise ValueError("パスワード付きPDFは読み込めません。") from exc
        if not unlocked:
            raise ValueError("パスワード付きPDFは読み込めません。")
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(data: bytes) -> str:
    _validate_office_archive(data)
    document = Document(BytesIO(data))
    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                chunks.append("\t".join(values))
    return "\n".join(chunks)


def _extract_xlsx_text(data: bytes) -> str:
    _validate_office_archive(data)
    workbook = load_workbook(BytesIO(data), read_only=True, data_only=False)
    chunks = []
    try:
        for sheet in workbook.worksheets:
            chunks.append(f"[シート: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    chunks.append("\t".join(values))
    finally:
        workbook.close()
    return "\n".join(chunks)


def _get_validated_image_mime_type(filename: str, data: bytes) -> str:
    extension = Path(filename).suffix.lower()
    mime_type = _AI_IMAGE_MIME_TYPES.get(extension)
    if not mime_type:
        raise ValueError("対応していない画像形式です。")
    if not data:
        raise ValueError("ファイルが空です。")
    if len(data) > _AI_FILE_MAX_BYTES:
        raise ValueError("ファイルサイズは10MB以下にしてください。")

    is_valid = False
    if extension == ".png":
        is_valid = data.startswith(b"\x89PNG\r\n\x1a\n")
    elif extension in {".jpg", ".jpeg"}:
        is_valid = data.startswith(b"\xff\xd8\xff")
    elif extension == ".webp":
        is_valid = len(data) >= 12 and data.startswith(b"RIFF") and data[8:12] == b"WEBP"
    else:
        heif_brands = {b"heic", b"heix", b"hevc", b"hevx", b"mif1", b"msf1"}
        is_valid = (
            len(data) >= 12
            and data[4:8] == b"ftyp"
            and data[8:12] in heif_brands
        )

    if not is_valid:
        raise ValueError("画像ファイルが破損しているか、形式が拡張子と一致しません。")
    return mime_type


def extract_ai_file_text(filename: str, data: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in _AI_DOCUMENT_EXTENSIONS:
        supported = ", ".join(sorted(_AI_FILE_EXTENSIONS))
        raise ValueError(f"対応していないファイル形式です。対応形式: {supported}")
    if not data:
        raise ValueError("ファイルが空です。")
    if len(data) > _AI_FILE_MAX_BYTES:
        raise ValueError("ファイルサイズは10MB以下にしてください。")

    try:
        if extension in {".txt", ".md", ".csv"}:
            text = _decode_text_file(data)
        elif extension == ".json":
            parsed = json.loads(_decode_text_file(data))
            text = json.dumps(parsed, ensure_ascii=False, indent=2)
        elif extension == ".pdf":
            text = _extract_pdf_text(data)
        elif extension == ".docx":
            text = _extract_docx_text(data)
        else:
            text = _extract_xlsx_text(data)
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("ファイルの内容を読み取れませんでした。") from exc

    text = text.replace("\x00", "").strip()
    if not text:
        raise ValueError("ファイルから文字を抽出できませんでした。画像のみのPDFには対応していません。")
    if len(text) > _AI_FILE_TEXT_MAX_CHARS:
        text = text[:_AI_FILE_TEXT_MAX_CHARS] + "\n\n[以降は文字数上限のため省略されました]"
    return text


def get_ai_input_from_request() -> tuple[str, genai_types.Part | None]:
    if request.is_json:
        data = request.get_json(silent=True) or {}
        prompt = str(data.get("prompt", "")).strip()
        upload = None
    else:
        prompt = str(request.form.get("prompt", "")).strip()
        upload = request.files.get("file")

    file_text = ""
    filename = ""
    image_part = None
    if upload and upload.filename:
        filename = upload.filename.replace("\\", "/").rsplit("/", 1)[-1][:255]
        file_data = upload.read(_AI_FILE_MAX_BYTES + 1)
        extension = Path(filename).suffix.lower()
        if extension in _AI_IMAGE_MIME_TYPES:
            mime_type = _get_validated_image_mime_type(filename, file_data)
            image_part = genai_types.Part.from_bytes(data=file_data, mime_type=mime_type)
        else:
            file_text = extract_ai_file_text(filename, file_data)

    if not prompt and not file_text and image_part is None:
        raise ValueError("テーマを入力するか、ファイルを選択してください。")
    if image_part is not None:
        instruction = prompt or "添付画像の内容を読み取り、重要な情報を整理してください。"
        return f"依頼・テーマ:\n{instruction}\n\n添付画像: {filename}", image_part
    if not file_text:
        return prompt, None

    instruction = prompt or "添付ファイルの内容を整理してください。"
    return (
        f"依頼・テーマ:\n{instruction}\n\n"
        f"添付ファイル「{filename}」の内容:\n"
        "--- ファイル内容ここから ---\n"
        f"{file_text}\n"
        "--- ファイル内容ここまで ---"
    ), None


def build_ai_contents(prompt_template: str, prompt: str, image_part: genai_types.Part | None):
    formatted_prompt = prompt_template.format(prompt=prompt)
    if image_part is None:
        return formatted_prompt
    return [image_part, formatted_prompt]


def _extract_mindmap_memo(node: dict, title: str, depth: int) -> str:
    for key in ("memo", "content", "description", "summary", "detail", "details"):
        value = node.get(key)
        if isinstance(value, list):
            text = "\n".join(str(item).strip() for item in value if str(item).strip())
        else:
            text = str(value or "").strip()
        if text:
            return text

    if depth == 0:
        return f"{title}の全体像、主要な観点、深掘りすべきポイントを整理します。"
    return f"{title}について、目的・要点・具体例を整理するためのメモです。"


def normalize_ai_mindmap_tree(tree):
    if not isinstance(tree, dict):
        raise ValueError("AIの出力形式が正しくありません。")

    def normalize_node(node, depth=0):
        if not isinstance(node, dict):
            return None
        title = str(node.get("title") or "トピック").strip()[:80] or "トピック"
        children = node.get("children")
        if not isinstance(children, list):
            children = []
        return {
            "title": title,
            "memo": _extract_mindmap_memo(node, title, depth),
            "children": [
                child
                for child in (normalize_node(child_node, depth + 1) for child_node in children)
                if child
            ],
        }

    return normalize_node(tree)


def _get_bearer_token() -> str | None:
    authorization = request.headers.get("Authorization", "")
    scheme, separator, token = authorization.partition(" ")
    if not separator or scheme.lower() != "bearer":
        return None
    token = token.strip()
    return token or None


_API_TOKEN_TYPE_READ = "read"
_API_TOKEN_TYPE_WRITE = "write"


def _lookup_user_token_uid(token_hash: str, token_type: str) -> str | None:
    """一般ユーザー用トークンのハッシュから、対応するuidをFirestoreの api_tokens コレクションで引く。"""
    doc = get_firestore_client().collection("api_tokens").document(token_hash).get()
    if not doc.exists:
        return None
    data = doc.to_dict() or {}
    if data.get("type") != token_type:
        return None
    uid = data.get("uid")
    return str(uid) if uid else None


def _is_valid_admin_token(token_hash_env: str, presented_hash: str) -> str | None:
    """管理者用の固定トークン（環境変数）と一致すればuidを返す。"""
    expected_hash = os.getenv(token_hash_env, "").strip().lower()
    admin_uid = os.getenv(_TODO_API_UID_ENV, "").strip()
    if (
        len(expected_hash) != 64
        or any(character not in "0123456789abcdef" for character in expected_hash)
        or not admin_uid
    ):
        return None
    if not hmac.compare_digest(presented_hash, expected_hash):
        return None
    return admin_uid


def _authenticate_todo_api_request(token_hash_env: str, token_type: str):
    token = _get_bearer_token()
    if token is None:
        return None, (
            jsonify(error="Bearerトークンが必要です。"),
            401,
            {
                "Cache-Control": "no-store",
                "WWW-Authenticate": 'Bearer realm="matometokiya-todos"',
            },
        )

    presented_hash = hashlib.sha256(token.encode("utf-8")).hexdigest()

    admin_uid = _is_valid_admin_token(token_hash_env, presented_hash)
    if admin_uid:
        return admin_uid, None

    try:
        user_uid = _lookup_user_token_uid(presented_hash, token_type)
    except Exception:
        app.logger.exception("Todo APIのトークン確認に失敗しました。")
        return None, (
            jsonify(error="トークンの確認に失敗しました。"),
            502,
            {"Cache-Control": "no-store"},
        )

    if user_uid:
        return user_uid, None

    return None, (
        jsonify(error="Bearerトークンが正しくありません。"),
        401,
        {
            "Cache-Control": "no-store",
            "WWW-Authenticate": 'Bearer realm="matometokiya-todos"',
        },
    )


def authenticate_todo_api_request():
    return _authenticate_todo_api_request(_TODO_API_TOKEN_HASH_ENV, _API_TOKEN_TYPE_READ)


def authenticate_todo_write_api_request():
    """完了マーク用の書き込みAPI認証。読み取り用とは別のトークンを要求する。"""
    return _authenticate_todo_api_request(_TODO_API_WRITE_TOKEN_HASH_ENV, _API_TOKEN_TYPE_WRITE)


def _ensure_firebase_app():
    """firebase_adminのデフォルトAppを、正しいプロジェクトIDで初期化しておく。"""
    import firebase_admin
    from firebase_admin import credentials

    # GOOGLE_CLOUD_PROJECT はVertex AI用に別プロジェクトを指す場合があるため、
    # Firestoreの接続先プロジェクトはFIREBASE_AUTH_DOMAINから明示的に固定する。
    firebase_project_id = FIREBASE_AUTH_DOMAIN.split(".")[0]
    options = {"projectId": firebase_project_id}

    try:
        firebase_admin.get_app()
    except ValueError:
        service_account_json = os.getenv("FIREBASE_SERVICE_ACCOUNT_JSON", "").strip()
        if service_account_json:
            credential = credentials.Certificate(json.loads(service_account_json))
            firebase_admin.initialize_app(credential, options)
        else:
            firebase_admin.initialize_app(options=options)


def get_firestore_client():
    global _firestore_client
    if _firestore_client is not None:
        return _firestore_client

    from firebase_admin import firestore

    _ensure_firebase_app()
    _firestore_client = firestore.client()
    return _firestore_client


def verify_firebase_id_token():
    """AuthorizationヘッダーのFirebase IDトークン（ログイン中の一般ユーザー用）を検証し、uidを返す。"""
    token = _get_bearer_token()
    if token is None:
        return None, (
            jsonify(error="ログイン情報が必要です。"),
            401,
            {"Cache-Control": "no-store"},
        )

    from firebase_admin import auth as firebase_auth

    _ensure_firebase_app()

    try:
        decoded = firebase_auth.verify_id_token(token)
    except Exception:
        return None, (
            jsonify(error="ログイン情報の確認に失敗しました。再度ログインし直してください。"),
            401,
            {"Cache-Control": "no-store"},
        )

    uid = decoded.get("uid")
    if not uid:
        return None, (
            jsonify(error="ログイン情報の確認に失敗しました。再度ログインし直してください。"),
            401,
            {"Cache-Control": "no-store"},
        )
    return uid, None


def _serialize_firestore_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat()
    return str(value)


def _project_name_for_note(note_id: str, notes_by_id: dict[str, dict]) -> str | None:
    current_id = note_id
    seen: set[str] = set()
    while current_id:
        if current_id in seen:
            return None
        seen.add(current_id)
        current_note = notes_by_id.get(current_id)
        if current_note is None:
            return None
        parent_id = current_note.get("parent_id")
        if parent_id is None:
            title = str(current_note.get("title") or "").strip()
            return title or None
        current_id = str(parent_id)
    return None


def _ancestor_titles(note_id: str, notes_by_id: dict[str, dict]) -> list[str]:
    """指定メモの祖先のタイトルを、ルート側から順に並べて返す（メモ自身は含まない）。
    例：制作 > RakuYade > homepage のように、どのアプリの話かを辿れるようにする。"""
    titles = []
    current_id = note_id
    seen: set[str] = set()
    while True:
        if current_id in seen:
            break
        seen.add(current_id)
        current_note = notes_by_id.get(current_id)
        if current_note is None:
            break
        parent_id = current_note.get("parent_id")
        if parent_id is None:
            break
        parent_id = str(parent_id)
        parent_note = notes_by_id.get(parent_id)
        if parent_note is None:
            break
        title = str(parent_note.get("title") or "").strip()
        if title:
            titles.append(title)
        current_id = parent_id

    titles.reverse()
    return titles


def _child_notes(note_id: str, notes_by_id: dict[str, dict]) -> list[dict]:
    """指定メモの直下の子メモを、order順に並べて返す。"""
    children = []
    for child_id, child in notes_by_id.items():
        if str(child.get("parent_id") or "") != note_id:
            continue
        order = child.get("order")
        order = order if isinstance(order, (int, float)) else 0
        title = str(child.get("title") or "").strip()
        content = str(child.get("content") or "").strip()
        if not title and not content:
            continue
        children.append({"order": order, "id": child_id, "title": title, "content": content})

    children.sort(key=lambda item: (item["order"], item["id"]))
    return [{"title": item["title"], "content": item["content"]} for item in children]


def build_incomplete_todo_payload(
    todo_documents: list[tuple[str, dict]],
    note_documents: list[tuple[str, dict]],
    project: str | None = None,
) -> list[dict]:
    notes_by_id = {document_id: data for document_id, data in note_documents}
    requested_project = project.strip() if project else None
    todos = []

    for document_id, data in todo_documents:
        if data.get("done") is not False:
            continue
        note_id = str(data.get("note_id") or "")
        note = notes_by_id.get(note_id)
        project_name = _project_name_for_note(note_id, notes_by_id) if note_id else None
        if requested_project is not None and project_name != requested_project:
            continue
        content = str((note or {}).get("title") or data.get("title") or "無題").strip()
        note_content = str((note or {}).get("content") or "").strip()
        children = _child_notes(note_id, notes_by_id) if note_id else []
        path = _ancestor_titles(note_id, notes_by_id) if note_id else []
        todos.append(
            {
                "id": document_id,
                "content": content or "無題",
                "note_content": note_content,
                "children": children,
                "path": path,
                "priority": None,
                "project": project_name,
                "created_at": _serialize_firestore_value(data.get("created_at")),
            }
        )

    return sorted(
        todos,
        key=lambda todo: (todo["created_at"], todo["id"]),
        reverse=True,
    )


def fetch_incomplete_todos_for_uid(uid: str, project: str | None = None) -> list[dict]:
    from google.cloud.firestore_v1.base_query import FieldFilter

    user_ref = get_firestore_client().collection("users").document(uid)
    todo_snapshot = (
        user_ref.collection("todos")
        .where(filter=FieldFilter("done", "==", False))
        .stream()
    )
    todo_documents = [(document.id, document.to_dict() or {}) for document in todo_snapshot]
    if not todo_documents:
        return []

    note_documents = [
        (document.id, data)
        for document in user_ref.collection("notes").stream()
        if not (data := document.to_dict() or {}).get("deleted")
    ]
    return build_incomplete_todo_payload(todo_documents, note_documents, project)


def build_all_notes_payload(
    note_documents: list[tuple[str, dict]], include_locked: bool = False
) -> list[dict]:
    """鍵付きメモ（`locked`）は、この利用者が「Claude連携」画面のトグルで許可している
    場合（include_locked=True）だけ本文を含める。既定はタイトルのみ。"""
    notes_by_id = {document_id: data for document_id, data in note_documents}

    entries = []
    for document_id, data in note_documents:
        title = str(data.get("title") or "").strip()
        content = str(data.get("content") or "").strip()
        if not title and not content:
            continue
        locked = bool(data.get("locked"))
        if locked and not include_locked:
            content = ""
        order = data.get("order")
        order = order if isinstance(order, (int, float)) else 0
        path = _ancestor_titles(document_id, notes_by_id)
        entries.append(
            (
                (len(path), path, order, document_id),
                {
                    "id": document_id,
                    "title": title or "無題",
                    "content": content,
                    "path": path,
                    "locked": locked,
                    "checked": bool(data.get("checked")),
                    "created_at": _serialize_firestore_value(data.get("created_at")),
                },
            )
        )

    entries.sort(key=lambda entry: entry[0])
    return [note for _sort_key, note in entries]


def fetch_all_notes_for_uid(uid: str, include_locked: bool = False) -> list[dict]:
    """このユーザーの全メモを、階層の浅い順に並べて返す（要約などClaude Code側での
    全量把握のための読み取り専用API向け。Todo取得APIと同じ読み取り用トークンを使う）。"""
    user_ref = get_firestore_client().collection("users").document(uid)
    note_documents = [
        (document.id, data)
        for document in user_ref.collection("notes").stream()
        if not (data := document.to_dict() or {}).get("deleted")
    ]
    return build_all_notes_payload(note_documents, include_locked)


def _note_ref(uid: str, note_id: str):
    return get_firestore_client().collection("users").document(uid).collection("notes").document(note_id)


def _next_order_for_new_note(uid: str, parent_id: str | None) -> int:
    """アプリ画面のnextOrderForNewNote()と同じ規則（末尾の兄弟 + 1000）で並び順を決める。"""
    notes_ref = get_firestore_client().collection("users").document(uid).collection("notes")
    max_order = None
    for doc in notes_ref.stream():
        data = doc.to_dict() or {}
        if data.get("deleted") or data.get("parent_id") != parent_id:
            continue
        order = data.get("order")
        if isinstance(order, (int, float)) and (max_order is None or order > max_order):
            max_order = order
    return int(max_order) + 1000 if max_order is not None else 1000


def create_note_for_uid(
    uid: str, title: str, content: str, parent_id: str | None
) -> dict | None:
    """新しいメモを作成する。parent_idを指定した場合、そのメモが存在し削除もされて
    いなければ子として作成する。parent_idが見つからなければNoneを返す
    （呼び出し元は404扱いにする）。"""
    if parent_id:
        parent_doc = _note_ref(uid, parent_id).get()
        if not parent_doc.exists or (parent_doc.to_dict() or {}).get("deleted"):
            return None

    note_id = secrets.token_hex(6)
    now = datetime.now(timezone.utc).isoformat()
    note = {
        "id": note_id,
        "parent_id": parent_id,
        "title": (title or "").strip()[:120] or "無題",
        "content": content or "",
        "created_at": now,
        "updated_at": now,
        "source_file": None,
        "media": [],
        "pinned": False,
        "checked": False,
        "checked_at": None,
        "locked": False,
        "order": _next_order_for_new_note(uid, parent_id),
    }
    _note_ref(uid, note_id).set(note)
    return note


def update_note_for_uid(
    uid: str,
    note_id: str,
    title: str | None,
    content: str | None,
    checked: bool | None = None,
) -> dict | None:
    """指定メモのtitle/content/checkedを更新する。存在しない、または削除済みならNoneを返す。"""
    ref = _note_ref(uid, note_id)
    doc = ref.get()
    if not doc.exists:
        return None
    data = doc.to_dict() or {}
    if data.get("deleted"):
        return None

    updates: dict = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if title is not None:
        updates["title"] = (title or "").strip()[:120] or "無題"
    if content is not None:
        updates["content"] = content or ""
    if checked is not None:
        updates["checked"] = bool(checked)
        updates["checked_at"] = datetime.now(timezone.utc).isoformat() if checked else None
    ref.update(updates)
    data.update(updates)
    data["id"] = note_id
    return data


def _collect_note_subtree_ids(uid: str, note_id: str) -> list[str]:
    """指定メモとその子孫（削除済みを除く）のIDを全て集める。"""
    notes_ref = get_firestore_client().collection("users").document(uid).collection("notes")
    children_by_parent: dict[str | None, list[str]] = {}
    for doc in notes_ref.stream():
        data = doc.to_dict() or {}
        if data.get("deleted"):
            continue
        children_by_parent.setdefault(data.get("parent_id"), []).append(doc.id)

    ids = []
    stack = [note_id]
    while stack:
        current = stack.pop()
        ids.append(current)
        stack.extend(children_by_parent.get(current, []))
    return ids


def delete_note_for_uid(uid: str, note_id: str) -> bool:
    """指定メモとその子孫メモを、アプリのゴミ箱と同じ形（deleted=Trueを立てるだけで
    Firestoreのドキュメント自体は消さない）で削除する。アプリ画面のゴミ箱から復元
    できる。存在しない、または既に削除済みなら何もせずFalseを返す。"""
    doc = _note_ref(uid, note_id).get()
    if not doc.exists or (doc.to_dict() or {}).get("deleted"):
        return False

    ids = _collect_note_subtree_ids(uid, note_id)
    db = get_firestore_client()
    notes_ref = db.collection("users").document(uid).collection("notes")
    now = datetime.now(timezone.utc).isoformat()
    batch = db.batch()
    for doc_id in ids:
        batch.update(notes_ref.document(doc_id), {"deleted": True, "deleted_at": now})
    batch.commit()
    return True


def mark_todo_done_for_uid(uid: str, todo_id: str) -> bool:
    """指定Todoをdone=Trueにする。対象が存在しなければFalseを返す。
    このTodoの元になっているメモ自体にも、アプリの画面上でチェック済みだと
    分かるようchecked/checked_atを立てる（Todo側だけdoneにしても、メモの
    チェックマークは連動しておらず利用者から見て完了したように見えない
    ため）。"""
    db = get_firestore_client()
    todo_ref = (
        db.collection("users")
        .document(uid)
        .collection("todos")
        .document(todo_id)
    )
    todo_doc = todo_ref.get()
    if not todo_doc.exists:
        return False
    todo_ref.update({"done": True})

    note_id = (todo_doc.to_dict() or {}).get("note_id")
    if note_id:
        note_ref = db.collection("users").document(uid).collection("notes").document(str(note_id))
        if note_ref.get().exists:
            note_ref.update({
                "checked": True,
                "checked_at": datetime.now(timezone.utc).isoformat(),
            })
    return True


def _generate_pairing_code() -> str:
    return f"{secrets.randbelow(10 ** _CLAUDE_PAIRING_CODE_LENGTH):0{_CLAUDE_PAIRING_CODE_LENGTH}d}"


def _revoke_user_tokens(uid: str) -> None:
    """このユーザーのClaude連携用トークンをすべて失効させる。"""
    from firebase_admin import firestore

    db = get_firestore_client()
    user_ref = db.collection("users").document(uid)
    user_doc = user_ref.get()
    data = user_doc.to_dict() or {} if user_doc.exists else {}

    for field in ("claude_read_token_hash", "claude_write_token_hash"):
        token_hash = data.get(field)
        if token_hash:
            db.collection("api_tokens").document(token_hash).delete()

    if user_doc.exists:
        user_ref.update(
            {
                "claude_connected_at": firestore.DELETE_FIELD,
                "claude_read_token_hash": firestore.DELETE_FIELD,
                "claude_write_token_hash": firestore.DELETE_FIELD,
            }
        )


def start_claude_connect(uid: str) -> dict:
    """短い引き換えコードだけを発行する。実際のAPIトークンの発行・登録や
    「連携済み」への反映は、コードが実際に引き換えられた時点
    （exchange_pairing_code）まで行わない。コードを生成しただけの時点で
    connected扱いにしてしまうと、Claude Code側にまだコードを渡していない
    のに画面上は「連携済み」と表示され、利用者が誤って古いトークンのまま
    連携完了だと勘違いする（実際には未引き換え＝旧トークンは失効済み）
    という不具合が起きるため。"""
    db = get_firestore_client()

    read_token = secrets.token_urlsafe(32)
    write_token = secrets.token_urlsafe(32)
    now = datetime.now(timezone.utc)

    code = _generate_pairing_code()
    for _ in range(5):
        if not db.collection("pairing_codes").document(code).get().exists:
            break
        code = _generate_pairing_code()
    else:
        raise RuntimeError("ペアリングコードの発行に失敗しました。")

    db.collection("pairing_codes").document(code).set(
        {
            "uid": uid,
            "read_token": read_token,
            "write_token": write_token,
            "expires_at": now + _CLAUDE_PAIRING_CODE_TTL,
        }
    )

    return {"code": code, "expires_in": int(_CLAUDE_PAIRING_CODE_TTL.total_seconds())}


def exchange_pairing_code(code: str) -> dict | None:
    """コードを使い捨てで引き換え、実際にAPIトークンを発行・登録してから返す。
    無効・期限切れならNoneを返す。"""
    db = get_firestore_client()
    doc_ref = db.collection("pairing_codes").document(code)
    doc = doc_ref.get()
    if not doc.exists:
        return None

    data = doc.to_dict() or {}
    doc_ref.delete()  # 成功・失敗にかかわらず、コードは1回きりで使い捨てる

    expires_at = data.get("expires_at")
    if not isinstance(expires_at, datetime):
        return None
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        return None

    uid = data.get("uid")
    read_token = data.get("read_token")
    write_token = data.get("write_token")
    if not uid or not read_token or not write_token:
        return None

    # このコードが実際に引き換えられた今の時点で初めて、トークンを有効化し
    # 「連携済み」にする。古いトークンがあればここで失効させる
    # （1ユーザー1組のトークンにする）。
    _revoke_user_tokens(uid)

    read_hash = hashlib.sha256(read_token.encode("utf-8")).hexdigest()
    write_hash = hashlib.sha256(write_token.encode("utf-8")).hexdigest()
    now = datetime.now(timezone.utc)

    db.collection("api_tokens").document(read_hash).set(
        {"uid": uid, "type": _API_TOKEN_TYPE_READ, "created_at": now}
    )
    db.collection("api_tokens").document(write_hash).set(
        {"uid": uid, "type": _API_TOKEN_TYPE_WRITE, "created_at": now}
    )
    db.collection("users").document(uid).set(
        {
            "claude_connected_at": now,
            "claude_read_token_hash": read_hash,
            "claude_write_token_hash": write_hash,
        },
        merge=True,
    )

    return {"read_token": read_token, "write_token": write_token}


def get_claude_connect_status(uid: str) -> str | None:
    doc = get_firestore_client().collection("users").document(uid).get()
    if not doc.exists:
        return None
    connected_at = (doc.to_dict() or {}).get("claude_connected_at")
    return _serialize_firestore_value(connected_at) if connected_at else None


def get_claude_include_locked_notes(uid: str) -> bool:
    """「Claude連携」画面のトグルで、この利用者が鍵付きメモの本文もClaudeに渡す設定に
    しているかどうか。既定はFalse（渡さない）。"""
    doc = get_firestore_client().collection("users").document(uid).get()
    if not doc.exists:
        return False
    return bool((doc.to_dict() or {}).get("claude_include_locked_notes"))


def set_claude_include_locked_notes(uid: str, include_locked: bool) -> None:
    get_firestore_client().collection("users").document(uid).set(
        {"claude_include_locked_notes": bool(include_locked)}, merge=True
    )


@app.get("/app")
def index():
    response = make_response(render_template("index.html"))
    response.headers["Cache-Control"] = "no-store, max-age=0"
    return response


@app.get("/auth/action")
def auth_action():
    response = make_response(render_template("auth_action.html"))
    response.headers["Cache-Control"] = "no-store, max-age=0"
    return response


def proxy_firebase_helper(path: str) -> Response:
    query = request.query_string.decode("utf-8")
    target_url = f"https://{FIREBASE_AUTH_DOMAIN}/{path}"
    if query:
        target_url = f"{target_url}?{query}"

    headers = {
        key: value
        for key, value in request.headers.items()
        if key.lower() not in {"host", "content-length", "connection", "accept-encoding"}
    }
    data = request.get_data() if request.method not in {"GET", "HEAD"} else None
    upstream_request = urllib.request.Request(
        target_url,
        data=data,
        headers=headers,
        method=request.method,
    )

    try:
        with urllib.request.urlopen(upstream_request, timeout=15) as upstream:
            body = upstream.read()
            response_headers = {
                key: value
                for key, value in upstream.headers.items()
                if key.lower() not in {"connection", "transfer-encoding", "content-encoding", "content-length"}
            }
            return Response(body, status=upstream.status, headers=response_headers)
    except urllib.error.HTTPError as error:
        body = error.read()
        response_headers = {
            key: value
            for key, value in error.headers.items()
            if key.lower() not in {"connection", "transfer-encoding", "content-encoding", "content-length"}
        }
        return Response(body, status=error.code, headers=response_headers)
    except urllib.error.URLError as error:
        return Response(f"Firebase auth helper proxy failed: {error.reason}", status=502)


@app.route("/__/auth/<path:auth_path>", methods=["GET", "POST", "HEAD", "OPTIONS"])
def firebase_auth_helper(auth_path: str):
    return proxy_firebase_helper(f"__/auth/{auth_path}")


@app.route("/__/firebase/init.json", methods=["GET", "HEAD"])
def firebase_init_json():
    return proxy_firebase_helper("__/firebase/init.json")


# ── SEOページ ─────────────────────────────────────────────────────────────────

@app.get("/")
def seo_landing():
    webapp_jsonld = {
        "@context": "https://schema.org",
        "@type": "WebApplication",
        "name": SITE_NAME,
        "url": f"{SITE_URL}/app",
        "description": (
            "頭の中のアイデアやタスクを、階層型のメモとマインドマップで整理できる無料アプリ。"
            "AIによる自動生成や複数人での共同編集にも対応しています。"
        ),
        "applicationCategory": "ProductivityApplication",
        "operatingSystem": "Web",
        "offers": {
            "@type": "Offer",
            "price": "0",
            "priceCurrency": "JPY",
        },
    }
    return render_template("seo/landing.html", webapp_jsonld=webapp_jsonld, faq_items=FAQ_ITEMS[:3])


@app.get("/how-to-use")
def how_to_use():
    return render_template("seo/how_to_use.html")


@app.get("/faq")
def faq():
    faq_jsonld = {
        "@context": "https://schema.org",
        "@type": "FAQPage",
        "mainEntity": [
            {
                "@type": "Question",
                "name": item["question"],
                "acceptedAnswer": {"@type": "Answer", "text": item["answer"]},
            }
            for item in FAQ_ITEMS
        ],
    }
    return render_template("seo/faq.html", faq_items=FAQ_ITEMS, faq_jsonld=faq_jsonld)


@app.get("/about")
def about():
    return render_template("seo/about.html")


@app.get("/privacy")
def privacy():
    return render_template("seo/privacy.html")


@app.get("/terms")
def terms():
    return render_template("seo/terms.html")


@app.get("/contact")
def contact():
    return render_template("seo/contact.html")


@app.get("/robots.txt")
def robots_txt():
    body = "\n".join([
        "User-agent: *",
        "Allow: /",
        f"Sitemap: {SITE_URL}/sitemap.xml",
        "",
    ])
    return Response(body, mimetype="text/plain")


@app.get("/sitemap.xml")
def sitemap_xml():
    entries = []
    for page in get_public_pages():
        template_path = TEMPLATES_DIR / page["template"]
        try:
            mtime = template_path.stat().st_mtime
        except OSError:
            mtime = datetime.now(tz=timezone.utc).timestamp()
        lastmod = datetime.fromtimestamp(mtime, tz=timezone.utc).strftime("%Y-%m-%d")
        entries.append(
            "<url>"
            f"<loc>{SITE_URL}{page['path']}</loc>"
            f"<lastmod>{lastmod}</lastmod>"
            f"<priority>{page['priority']}</priority>"
            "</url>"
        )
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">'
        + "".join(entries)
        + "</urlset>"
    )
    return Response(xml, mimetype="application/xml")


@app.get("/blog/<slug>")
def blog_post(slug):
    post = BLOG_POSTS_BY_SLUG.get(slug)
    if not post:
        abort(404)
    return render_template(post["template"])


@app.get("/<slug>")
def keyword_page(slug):
    page = KEYWORD_PAGES_BY_SLUG.get(slug)
    if not page:
        abort(404)
    return render_template(page["template"], keyword=page["keyword"])


@app.errorhandler(413)
def request_too_large(_error):
    return jsonify(error="ファイルサイズは10MB以下にしてください。"), 413


@app.get("/api/v1/todos")
def api_incomplete_todos():
    uid, auth_error = authenticate_todo_api_request()
    if auth_error is not None:
        return auth_error

    project = request.args.get("project")
    if project is not None:
        project = project.strip()
        if not project:
            return jsonify(error="projectは空にできません。"), 400, {"Cache-Control": "no-store"}
        if len(project) > 120:
            return jsonify(error="projectは120文字以内で指定してください。"), 400, {"Cache-Control": "no-store"}

    try:
        todos = fetch_incomplete_todos_for_uid(uid, project)
    except Exception:
        app.logger.exception("Todo APIでFirestoreの読み取りに失敗しました。")
        return (
            jsonify(error="Todoの取得に失敗しました。"),
            502,
            {"Cache-Control": "no-store"},
        )

    response = jsonify(
        todos=todos,
        count=len(todos),
        project=project,
        read_only=True,
    )
    response.headers["Cache-Control"] = "no-store"
    return response


@app.get("/api/v1/notes")
def api_all_notes():
    """全メモの一覧を返す（読み取り専用）。Todo取得APIと同じ読み取り用トークンで認証する。
    鍵付きメモの本文は、この利用者が「Claude連携」画面のトグルで許可している場合だけ含める。"""
    uid, auth_error = authenticate_todo_api_request()
    if auth_error is not None:
        return auth_error

    try:
        include_locked = get_claude_include_locked_notes(uid)
        notes = fetch_all_notes_for_uid(uid, include_locked)
    except Exception:
        app.logger.exception("Notes APIでFirestoreの読み取りに失敗しました。")
        return (
            jsonify(error="メモの取得に失敗しました。"),
            502,
            {"Cache-Control": "no-store"},
        )

    response = jsonify(notes=notes, count=len(notes), read_only=True)
    response.headers["Cache-Control"] = "no-store"
    return response


@app.post("/api/v1/notes")
def api_create_note():
    """新しいメモを作成する（書き込み専用トークン）。"""
    uid, auth_error = authenticate_todo_write_api_request()
    if auth_error is not None:
        return auth_error

    payload = request.get_json(silent=True) or {}
    title = str(payload.get("title") or "")
    content = str(payload.get("content") or "")
    parent_id = payload.get("parent_id")
    parent_id = str(parent_id) if parent_id else None

    try:
        note = create_note_for_uid(uid, title, content, parent_id)
    except Exception:
        app.logger.exception("Notes APIでメモの作成に失敗しました。")
        return jsonify(error="メモの作成に失敗しました。"), 502, {"Cache-Control": "no-store"}

    if note is None:
        return (
            jsonify(error="指定されたparent_idのメモが見つかりません。"),
            404,
            {"Cache-Control": "no-store"},
        )

    response = jsonify(id=note["id"], title=note["title"], content=note["content"], parent_id=note["parent_id"])
    response.headers["Cache-Control"] = "no-store"
    return response, 201


@app.patch("/api/v1/notes/<note_id>")
def api_update_note(note_id):
    """既存メモのtitle/content/checkedを更新する（書き込み専用トークン）。
    鍵付きメモは、この利用者が「Claude連携」画面のトグルで許可していない限り更新できない
    （読み取れないメモをClaudeが書き換えられてしまうのを防ぐため）。"""
    uid, auth_error = authenticate_todo_write_api_request()
    if auth_error is not None:
        return auth_error

    note_id = (note_id or "").strip()
    if not note_id or len(note_id) > 200:
        return jsonify(error="note_idが正しくありません。"), 400, {"Cache-Control": "no-store"}

    payload = request.get_json(silent=True) or {}
    if "title" not in payload and "content" not in payload and "checked" not in payload:
        return (
            jsonify(error="title、content、checkedのいずれかを指定してください。"),
            400,
            {"Cache-Control": "no-store"},
        )
    title = str(payload["title"]) if "title" in payload else None
    content = str(payload["content"]) if "content" in payload else None
    checked = bool(payload["checked"]) if "checked" in payload else None

    try:
        existing = _note_ref(uid, note_id).get()
        if existing.exists and (existing.to_dict() or {}).get("locked") and not get_claude_include_locked_notes(uid):
            return (
                jsonify(error="鍵付きメモです。「Claude連携」設定で許可してから編集してください。"),
                403,
                {"Cache-Control": "no-store"},
            )
        note = update_note_for_uid(uid, note_id, title, content, checked)
    except Exception:
        app.logger.exception("Notes APIでメモの更新に失敗しました。")
        return jsonify(error="メモの更新に失敗しました。"), 502, {"Cache-Control": "no-store"}

    if note is None:
        return jsonify(error="指定されたメモが見つかりません。"), 404, {"Cache-Control": "no-store"}

    response = jsonify(id=note_id, title=note["title"], content=note["content"], checked=bool(note.get("checked")))
    response.headers["Cache-Control"] = "no-store"
    return response


@app.delete("/api/v1/notes/<note_id>")
def api_delete_note(note_id):
    """指定メモとその子孫メモを削除する（書き込み専用トークン）。アプリのゴミ箱と同じ
    ソフトデリートで、アプリ画面から復元できる。鍵付きメモは、この利用者が
    「Claude連携」画面のトグルで許可していない限り削除できない。"""
    uid, auth_error = authenticate_todo_write_api_request()
    if auth_error is not None:
        return auth_error

    note_id = (note_id or "").strip()
    if not note_id or len(note_id) > 200:
        return jsonify(error="note_idが正しくありません。"), 400, {"Cache-Control": "no-store"}

    try:
        existing = _note_ref(uid, note_id).get()
        if existing.exists and (existing.to_dict() or {}).get("locked") and not get_claude_include_locked_notes(uid):
            return (
                jsonify(error="鍵付きメモです。「Claude連携」設定で許可してから削除してください。"),
                403,
                {"Cache-Control": "no-store"},
            )
        deleted = delete_note_for_uid(uid, note_id)
    except Exception:
        app.logger.exception("Notes APIでメモの削除に失敗しました。")
        return jsonify(error="メモの削除に失敗しました。"), 502, {"Cache-Control": "no-store"}

    if not deleted:
        return jsonify(error="指定されたメモが見つかりません。"), 404, {"Cache-Control": "no-store"}

    response = jsonify(id=note_id, deleted=True)
    response.headers["Cache-Control"] = "no-store"
    return response


@app.get("/api/v1/claude-connect/locked-notes-setting")
def api_claude_connect_locked_notes_setting_get():
    """「鍵付きメモの内容もClaude連携で取得できるようにする」設定の現在値を返す。
    ログイン中の本人のみ（Firebase IDトークン認証）。"""
    uid, auth_error = verify_firebase_id_token()
    if auth_error is not None:
        return auth_error

    try:
        include_locked = get_claude_include_locked_notes(uid)
    except Exception:
        app.logger.exception("鍵付きメモ設定の取得に失敗しました。")
        return jsonify(error="設定の取得に失敗しました。"), 502, {"Cache-Control": "no-store"}

    response = jsonify(include_locked=include_locked)
    response.headers["Cache-Control"] = "no-store"
    return response


@app.post("/api/v1/claude-connect/locked-notes-setting")
def api_claude_connect_locked_notes_setting_post():
    """「鍵付きメモの内容もClaude連携で取得できるようにする」設定を変更する。
    ログイン中の本人のみ（Firebase IDトークン認証）。"""
    uid, auth_error = verify_firebase_id_token()
    if auth_error is not None:
        return auth_error

    payload = request.get_json(silent=True) or {}
    include_locked = bool(payload.get("include_locked"))

    try:
        set_claude_include_locked_notes(uid, include_locked)
    except Exception:
        app.logger.exception("鍵付きメモ設定の更新に失敗しました。")
        return jsonify(error="設定の更新に失敗しました。"), 502, {"Cache-Control": "no-store"}

    response = jsonify(include_locked=include_locked)
    response.headers["Cache-Control"] = "no-store"
    return response


@app.post("/api/v1/todos/<todo_id>/complete")
def api_complete_todo(todo_id):
    uid, auth_error = authenticate_todo_write_api_request()
    if auth_error is not None:
        return auth_error

    todo_id = (todo_id or "").strip()
    if not todo_id or len(todo_id) > 200:
        return jsonify(error="todo_idが正しくありません。"), 400, {"Cache-Control": "no-store"}

    try:
        found = mark_todo_done_for_uid(uid, todo_id)
    except Exception:
        app.logger.exception("Todo APIでFirestoreの書き込みに失敗しました。")
        return (
            jsonify(error="Todoの更新に失敗しました。"),
            502,
            {"Cache-Control": "no-store"},
        )

    if not found:
        return jsonify(error="指定されたTodoが見つかりません。"), 404, {"Cache-Control": "no-store"}

    response = jsonify(id=todo_id, done=True)
    response.headers["Cache-Control"] = "no-store"
    return response


@app.post("/api/v1/claude-connect/start")
def api_claude_connect_start():
    """ログイン中のユーザー用に、Claude Code連携のペアリングコードを発行する。"""
    uid, auth_error = verify_firebase_id_token()
    if auth_error is not None:
        return auth_error

    try:
        result = start_claude_connect(uid)
    except Exception:
        app.logger.exception("Claude連携の開始に失敗しました。")
        return (
            jsonify(error="連携の開始に失敗しました。時間をおいて再度お試しください。"),
            502,
            {"Cache-Control": "no-store"},
        )

    response = jsonify(**result)
    response.headers["Cache-Control"] = "no-store"
    return response


@app.post("/api/v1/claude-connect/exchange")
def api_claude_connect_exchange():
    """ペアリングコードを、実際のAPIトークンと引き換える（ログイン不要、コード自体が認証）。"""
    payload = request.get_json(silent=True) or {}
    code = str(payload.get("code") or "").strip()
    if not code.isdigit() or len(code) != _CLAUDE_PAIRING_CODE_LENGTH:
        return jsonify(error="コードが正しくありません。"), 400, {"Cache-Control": "no-store"}

    try:
        result = exchange_pairing_code(code)
    except Exception:
        app.logger.exception("ペアリングコードの引き換えに失敗しました。")
        return (
            jsonify(error="引き換えに失敗しました。時間をおいて再度お試しください。"),
            502,
            {"Cache-Control": "no-store"},
        )

    if result is None:
        return (
            jsonify(error="コードが無効か、有効期限が切れています。もう一度発行してください。"),
            404,
            {"Cache-Control": "no-store"},
        )

    response = jsonify(**result)
    response.headers["Cache-Control"] = "no-store"
    return response


@app.get("/api/v1/claude-connect/status")
def api_claude_connect_status():
    uid, auth_error = verify_firebase_id_token()
    if auth_error is not None:
        return auth_error

    try:
        connected_at = get_claude_connect_status(uid)
    except Exception:
        app.logger.exception("Claude連携状況の取得に失敗しました。")
        return jsonify(error="状況の取得に失敗しました。"), 502, {"Cache-Control": "no-store"}

    response = jsonify(connected=connected_at is not None, connected_at=connected_at)
    response.headers["Cache-Control"] = "no-store"
    return response


@app.post("/api/v1/claude-connect/revoke")
def api_claude_connect_revoke():
    uid, auth_error = verify_firebase_id_token()
    if auth_error is not None:
        return auth_error

    try:
        _revoke_user_tokens(uid)
    except Exception:
        app.logger.exception("Claude連携の解除に失敗しました。")
        return jsonify(error="連携の解除に失敗しました。"), 502, {"Cache-Control": "no-store"}

    response = jsonify(connected=False)
    response.headers["Cache-Control"] = "no-store"
    return response


@app.post("/api/ai-note")
def api_ai_note():
    try:
        prompt, image_part = get_ai_input_from_request()
    except ValueError as e:
        return jsonify(error=str(e)), 400

    try:
        client = create_gemini_client()
        response = client.models.generate_content(
            model=_GEMINI_MODEL,
            contents=build_ai_contents(_NOTE_PROMPT, prompt, image_part),
            config=genai_types.GenerateContentConfig(
                response_mime_type="application/json",
                max_output_tokens=4096,
                temperature=0.7,
            ),
        )
        raw = response.text.strip()
        if raw.startswith("```"):
            lines = raw.splitlines()
            raw = "\n".join(lines[1:-1] if lines[-1].startswith("```") else lines[1:])
        tree = json.loads(raw)
    except Exception as e:
        return jsonify(error=str(e)), 502

    return jsonify(tree=tree)


@app.post("/api/ai-mindmap")
def api_ai_mindmap():
    try:
        prompt, image_part = get_ai_input_from_request()
    except ValueError as e:
        return jsonify(error=str(e)), 400

    try:
        client = create_gemini_client()
        response = client.models.generate_content(
            model=_GEMINI_MODEL,
            contents=build_ai_contents(_MINDMAP_PROMPT, prompt, image_part),
            config=genai_types.GenerateContentConfig(
                response_mime_type="application/json",
                max_output_tokens=4096,
                temperature=0.7,
            ),
        )
        raw = response.text.strip()
        if raw.startswith("```"):
            lines = raw.splitlines()
            raw = "\n".join(lines[1:-1] if lines[-1].startswith("```") else lines[1:])
        tree = normalize_ai_mindmap_tree(json.loads(raw))
    except Exception as e:
        return jsonify(error=str(e)), 502

    return jsonify(tree=tree)


if __name__ == "__main__":
    port = int(os.getenv("PORT", "5006"))
    app.run(debug=True, port=port)
